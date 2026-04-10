#!/usr/bin/env python3
"""
OpenClaw Skill Gateway
======================
Lightweight HTTP API server that wraps Docker-based OpenClaw skills.
Runs on localhost:5050, exposes each skill as a POST endpoint.

No external dependencies -- stdlib only (http.server, subprocess, json).

Usage:
    python3 skill-gateway.py
    # or
    chmod +x skill-gateway.py && ./skill-gateway.py
"""

import base64
import json
import os
import shlex
import subprocess
import sys
import uuid
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse

HOST = "0.0.0.0"
PORT = 5050

SKILL_INPUT_DIR = "/tmp/skill-input"
SKILL_OUTPUT_DIR = "/tmp/skill-output"

SKILLS_BASE = os.path.dirname(os.path.abspath(__file__))

AVAILABLE_SKILLS = [
    "code-runner",
    "web-browser",
    "document-analyzer",
    "document-creator",
    "diagram-generator",
    "web-search",
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _log(method: str, path: str, status: int, detail: str = "") -> None:
    extra = f" | {detail}" if detail else ""
    print(f"[gateway] {method} {path} -> {status}{extra}", flush=True)


def _json_response(handler: BaseHTTPRequestHandler, data: dict, status: int = 200) -> None:
    body = json.dumps(data, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
    handler.send_header("Access-Control-Allow-Headers", "Content-Type")
    handler.end_headers()
    handler.wfile.write(body)


def _read_body(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", 0))
    if length == 0:
        return {}
    raw = handler.rfile.read(length)
    return json.loads(raw)


def _run_docker(args: list[str], timeout: int = 120) -> tuple[str, str, int]:
    """Run a Docker command and return (stdout, stderr, returncode)."""
    cmd = ["docker"] + args
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        return result.stdout, result.stderr, result.returncode
    except subprocess.TimeoutExpired:
        return "", f"Docker command timed out after {timeout}s", 124
    except FileNotFoundError:
        return "", "docker executable not found on PATH", 127


def _run_local(args, env=None, timeout=60):
    """Run a local command and return (stdout, stderr, returncode)."""
    merged_env = {**os.environ, **(env or {})}
    try:
        result = subprocess.run(
            args,
            capture_output=True,
            text=True,
            timeout=timeout,
            env=merged_env,
        )
        return result.stdout, result.stderr, result.returncode
    except subprocess.TimeoutExpired:
        return "", f"Command timed out after {timeout}s", 124
    except FileNotFoundError:
        return "", f"executable not found: {args[0]}", 127


def _parse_json_stdout(stdout: str) -> dict:
    """Try to parse JSON from stdout; fall back to wrapping raw text."""
    stdout = stdout.strip()
    if not stdout:
        return {"output": ""}
    try:
        return json.loads(stdout)
    except json.JSONDecodeError:
        return {"output": stdout}


# ---------------------------------------------------------------------------
# Skill handlers
# ---------------------------------------------------------------------------

def handle_code_runner_python(body: dict) -> tuple[dict, int]:
    code = body.get("code", "")
    timeout = body.get("timeout", 30)
    if not code:
        return {"error": "missing 'code' field"}, 400

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "openclaw-code-runner",
        "scripts/run_python.py", "-c", code,
        "--timeout", str(timeout),
    ], timeout=timeout + 30)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500
    return _parse_json_stdout(stdout), 200


def handle_code_runner_node(body: dict) -> tuple[dict, int]:
    code = body.get("code", "")
    timeout = body.get("timeout", 30)
    if not code:
        return {"error": "missing 'code' field"}, 400

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "--entrypoint", "python3",
        "openclaw-code-runner",
        "scripts/run_node.py", "-c", code,
        "--timeout", str(timeout),
    ], timeout=timeout + 30)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500
    return _parse_json_stdout(stdout), 200


def handle_web_browser_browse(body: dict) -> tuple[dict, int]:
    url = body.get("url", "")
    if not url:
        return {"error": "missing 'url' field"}, 400

    wait = body.get("wait", "load")
    timeout = body.get("timeout", 15000)

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "openclaw-web-browser",
        "scripts/browse.py", url,
        "--wait", str(wait),
        "--timeout", str(timeout),
    ], timeout=max(timeout // 1000, 30) + 30)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500
    return _parse_json_stdout(stdout), 200


def handle_web_browser_screenshot(body: dict) -> tuple[dict, int]:
    url = body.get("url", "")
    if not url:
        return {"error": "missing 'url' field"}, 400

    vw = body.get("viewport_width", 1280)
    vh = body.get("viewport_height", 720)

    # Ensure output dir is clean for this request
    request_dir = os.path.join(SKILL_OUTPUT_DIR, str(uuid.uuid4()))
    os.makedirs(request_dir, exist_ok=True)

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "-v", f"{request_dir}:/workspace",
        "openclaw-web-browser",
        "scripts/screenshot.py", url,
        "--viewport-width", str(vw),
        "--viewport-height", str(vh),
        "-o", "/workspace/screenshot.png",
    ], timeout=60)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500

    screenshot_path = os.path.join(request_dir, "screenshot.png")
    if not os.path.exists(screenshot_path):
        return {"error": "screenshot file was not produced by the container"}, 500

    with open(screenshot_path, "rb") as f:
        image_b64 = base64.b64encode(f.read()).decode("ascii")

    # Cleanup
    try:
        os.remove(screenshot_path)
        os.rmdir(request_dir)
    except OSError:
        pass

    return {"image": image_b64}, 200


def handle_web_browser_extract(body: dict) -> tuple[dict, int]:
    url = body.get("url", "")
    if not url:
        return {"error": "missing 'url' field"}, 400

    selector = body.get("selector")

    cmd = [
        "run", "--rm",
        "openclaw-web-browser",
        "scripts/extract_content.py", url,
    ]
    if selector:
        cmd.extend(["--selector", selector])

    stdout, stderr, rc = _run_docker(cmd, timeout=60)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500

    parsed = _parse_json_stdout(stdout)
    # Normalize to {"content": ...} if script returned raw text
    if "content" not in parsed and "output" in parsed:
        parsed = {"content": parsed["output"]}
    return parsed, 200


def handle_document_analyzer(body: dict) -> tuple[dict, int]:
    file_b64 = body.get("file_base64", "")
    filename = body.get("filename", "")
    if not file_b64 or not filename:
        return {"error": "missing 'file_base64' or 'filename'"}, 400

    # Sanitize filename to prevent path traversal
    safe_filename = os.path.basename(filename)
    if not safe_filename:
        return {"error": "invalid filename"}, 400

    request_dir = os.path.join(SKILL_INPUT_DIR, str(uuid.uuid4()))
    os.makedirs(request_dir, exist_ok=True)
    file_path = os.path.join(request_dir, safe_filename)

    try:
        with open(file_path, "wb") as f:
            f.write(base64.b64decode(file_b64))
    except Exception as exc:
        return {"error": f"failed to decode file_base64: {exc}"}, 400

    python_code = (
        "from pypdf import PdfReader; "
        f"reader = PdfReader('/input/{safe_filename}'); "
        "print('\\n'.join(p.extract_text() or '' for p in reader.pages))"
    )

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "-v", f"{request_dir}:/input",
        "openclaw-document-analyzer",
        "-c", python_code,
    ], timeout=60)

    # Cleanup
    try:
        os.remove(file_path)
        os.rmdir(request_dir)
    except OSError:
        pass

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500

    return {"text": stdout.strip()}, 200


def handle_document_creator(body: dict) -> tuple[dict, int]:
    doc_type = body.get("type", "docx")
    content = body.get("content", {})
    if not content:
        return {"error": "missing 'content' field"}, 400

    title = content.get("title", "Untitled")
    paragraphs = content.get("paragraphs", [])

    request_id = str(uuid.uuid4())
    output_dir = os.path.join(SKILL_OUTPUT_DIR, request_id)
    os.makedirs(output_dir, exist_ok=True)

    output_filename = f"output.{doc_type}"

    if doc_type == "docx":
        # Generate a docx using python-docx inside the container
        paragraphs_json = json.dumps(paragraphs)
        python_code = (
            "import json; "
            "from docx import Document; "
            "doc = Document(); "
            f"doc.add_heading({json.dumps(title)}, 0); "
            f"paras = json.loads({json.dumps(paragraphs_json)}); "
            "[doc.add_paragraph(p) for p in paras]; "
            f"doc.save('/output/{output_filename}')"
        )
    elif doc_type == "txt":
        text_content = f"{title}\n\n" + "\n\n".join(paragraphs)
        python_code = (
            f"with open('/output/{output_filename}', 'w') as f: "
            f"    f.write({json.dumps(text_content)})"
        )
    else:
        return {"error": f"unsupported document type: {doc_type}"}, 400

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "-v", f"{output_dir}:/output",
        "openclaw-document-creator",
        "-c", python_code,
    ], timeout=60)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500

    output_path = os.path.join(output_dir, output_filename)
    if not os.path.exists(output_path):
        return {"error": "output file was not produced by the container"}, 500

    with open(output_path, "rb") as f:
        file_b64 = base64.b64encode(f.read()).decode("ascii")

    # Cleanup
    try:
        os.remove(output_path)
        os.rmdir(output_dir)
    except OSError:
        pass

    return {"file_base64": file_b64, "filename": output_filename}, 200


def handle_diagram_graphviz(body: dict) -> tuple[dict, int]:
    dot = body.get("dot", "")
    fmt = body.get("format", "svg")
    if not dot:
        return {"error": "missing 'dot' field"}, 400
    if fmt not in ("svg", "png", "pdf"):
        return {"error": f"unsupported format: {fmt}. Use svg, png, or pdf"}, 400

    # Pipe the dot source into the `dot` command via bash -c
    # Use heredoc-style to avoid shell escaping issues with the dot content
    bash_script = f"echo {shlex.quote(dot)} | dot -T{fmt}"

    stdout, stderr, rc = _run_docker([
        "run", "--rm",
        "--platform", "linux/amd64",
        "--entrypoint", "bash",
        "openclaw-diagram-generator",
        "-c", bash_script,
    ], timeout=60)

    if rc != 0:
        return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500

    if fmt == "svg":
        return {"svg": stdout}, 200
    else:
        # Binary output -- base64-encode it
        # Re-run with raw bytes capture since text mode might corrupt binary
        result = subprocess.run(
            ["docker", "run", "--rm",
             "--platform", "linux/amd64",
             "--entrypoint", "bash",
             "openclaw-diagram-generator",
             "-c", bash_script],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            return {"error": result.stderr.decode(errors="replace").strip()}, 500
        return {"image": base64.b64encode(result.stdout).decode("ascii")}, 200


def handle_diagram_mermaid(body: dict) -> tuple[dict, int]:
    mermaid_src = body.get("mermaid", "")
    fmt = body.get("format", "svg")
    if not mermaid_src:
        return {"error": "missing 'mermaid' field"}, 400
    if fmt not in ("svg", "png", "pdf"):
        return {"error": f"unsupported format: {fmt}. Use svg, png, or pdf"}, 400

    escaped = shlex.quote(mermaid_src)
    bash_script = (
        f"echo {escaped} > /tmp/d.mmd && "
        f"python3 scripts/generate_mermaid.py /tmp/d.mmd -o /tmp/out.{fmt} && "
        f"cat /tmp/out.{fmt}"
    )

    if fmt == "svg":
        stdout, stderr, rc = _run_docker([
            "run", "--rm",
            "--platform", "linux/amd64",
            "--entrypoint", "bash",
            "openclaw-diagram-generator",
            "-c", bash_script,
        ], timeout=60)

        if rc != 0:
            return {"error": stderr.strip() or f"docker exited with code {rc}", "exit_code": rc}, 500
        return {"svg": stdout}, 200
    else:
        result = subprocess.run(
            ["docker", "run", "--rm",
             "--platform", "linux/amd64",
             "--entrypoint", "bash",
             "openclaw-diagram-generator",
             "-c", bash_script],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            return {"error": result.stderr.decode(errors="replace").strip()}, 500
        return {"image": base64.b64encode(result.stdout).decode("ascii")}, 200


def handle_web_search_searxng(body: dict) -> tuple[dict, int]:
    query = body.get("query", "")
    if not query:
        return {"error": "missing 'query' field"}, 400

    max_results = body.get("max_results", 5)
    categories = body.get("categories", "general")

    search_script = os.path.join(SKILLS_BASE, "web-search-searxng", "scripts", "search.py")
    if not os.path.exists(search_script):
        return {"error": f"search script not found: {search_script}"}, 500

    cmd = [
        "python3", search_script,
        query,
        "--max-results", str(max_results),
        "--categories", str(categories),
    ]
    env = {"SEARXNG_URL": "http://localhost:8888"}

    stdout, stderr, rc = _run_local(cmd, env=env, timeout=30)

    if rc != 0:
        return {"error": stderr.strip() or f"search exited with code {rc}", "exit_code": rc}, 500

    # Try to parse JSON array
    stdout = stdout.strip()
    try:
        results = json.loads(stdout)
        return {"results": results} if isinstance(results, list) else results, 200
    except json.JSONDecodeError:
        return {"results": stdout}, 200


# ---------------------------------------------------------------------------
# Route table
# ---------------------------------------------------------------------------

ROUTES = {
    "/skills/code-runner/python":       handle_code_runner_python,
    "/skills/code-runner/node":         handle_code_runner_node,
    "/skills/web-browser/browse":       handle_web_browser_browse,
    "/skills/web-browser/screenshot":   handle_web_browser_screenshot,
    "/skills/web-browser/extract":      handle_web_browser_extract,
    "/skills/document-analyzer/analyze": handle_document_analyzer,
    "/skills/document-creator/create":  handle_document_creator,
    "/skills/diagram-generator/graphviz": handle_diagram_graphviz,
    "/skills/diagram-generator/mermaid": handle_diagram_mermaid,
    "/skills/web-search/searxng":       handle_web_search_searxng,
}


# ---------------------------------------------------------------------------
# API documentation
# ---------------------------------------------------------------------------

API_DOCS = {
    "name": "OpenClaw Skill Gateway",
    "version": "1.0.0",
    "description": "HTTP API wrapping Docker-based OpenClaw skills",
    "base_url": f"http://localhost:{PORT}",
    "endpoints": [
        {
            "method": "GET",
            "path": "/",
            "description": "API documentation (this page)",
        },
        {
            "method": "GET",
            "path": "/health",
            "description": "Health check and list of available skills",
        },
        {
            "method": "POST",
            "path": "/skills/code-runner/python",
            "description": "Execute Python code in a sandboxed container",
            "body": {"code": "print('hello')", "timeout": 30},
        },
        {
            "method": "POST",
            "path": "/skills/code-runner/node",
            "description": "Execute Node.js code in a sandboxed container",
            "body": {"code": "console.log('hi')", "timeout": 30},
        },
        {
            "method": "POST",
            "path": "/skills/web-browser/browse",
            "description": "Browse a URL and return page content",
            "body": {"url": "https://example.com", "wait": "load", "timeout": 15000},
        },
        {
            "method": "POST",
            "path": "/skills/web-browser/screenshot",
            "description": "Take a screenshot of a URL",
            "body": {"url": "https://example.com", "viewport_width": 1280, "viewport_height": 720},
        },
        {
            "method": "POST",
            "path": "/skills/web-browser/extract",
            "description": "Extract content from a URL as markdown",
            "body": {"url": "https://example.com", "selector": "body"},
        },
        {
            "method": "POST",
            "path": "/skills/document-analyzer/analyze",
            "description": "Analyze a document (PDF). Send file as base64.",
            "body": {"file_base64": "<base64>", "filename": "doc.pdf"},
        },
        {
            "method": "POST",
            "path": "/skills/document-creator/create",
            "description": "Create a document (docx, txt)",
            "body": {"type": "docx", "content": {"title": "Title", "paragraphs": ["Para 1", "Para 2"]}},
        },
        {
            "method": "POST",
            "path": "/skills/diagram-generator/graphviz",
            "description": "Generate a diagram from Graphviz DOT source",
            "body": {"dot": "digraph { A -> B }", "format": "svg"},
        },
        {
            "method": "POST",
            "path": "/skills/diagram-generator/mermaid",
            "description": "Generate a diagram from Mermaid source",
            "body": {"mermaid": "graph TD; A-->B", "format": "svg"},
        },
        {
            "method": "POST",
            "path": "/skills/web-search/searxng",
            "description": "Search the web via local SearXNG instance",
            "body": {"query": "search term", "max_results": 5, "categories": "general"},
        },
    ],
}


# ---------------------------------------------------------------------------
# HTTP Handler
# ---------------------------------------------------------------------------

class SkillGatewayHandler(BaseHTTPRequestHandler):
    """Single-class request handler for all skill endpoints."""

    def log_message(self, format, *args):
        # Suppress default stderr logging; we use our own _log()
        pass

    def _send_cors_preflight(self) -> None:
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.send_header("Content-Length", "0")
        self.end_headers()

    def do_OPTIONS(self) -> None:
        self._send_cors_preflight()

    def do_GET(self) -> None:
        path = urlparse(self.path).path.rstrip("/") or "/"

        if path == "/":
            _log("GET", path, 200)
            _json_response(self, API_DOCS)
            return

        if path == "/health":
            _log("GET", path, 200)
            _json_response(self, {"status": "ok", "skills": AVAILABLE_SKILLS})
            return

        _log("GET", path, 404)
        _json_response(self, {"error": f"unknown endpoint: {path}"}, 404)

    def do_POST(self) -> None:
        path = urlparse(self.path).path.rstrip("/")
        handler_fn = ROUTES.get(path)

        if handler_fn is None:
            _log("POST", path, 404)
            _json_response(self, {"error": f"unknown endpoint: {path}"}, 404)
            return

        try:
            body = _read_body(self)
        except (json.JSONDecodeError, ValueError) as exc:
            _log("POST", path, 400, f"bad JSON: {exc}")
            _json_response(self, {"error": f"invalid JSON body: {exc}"}, 400)
            return

        try:
            result, status = handler_fn(body)
            _log("POST", path, status, handler_fn.__name__)
            _json_response(self, result, status)
        except Exception as exc:
            _log("POST", path, 500, f"unhandled: {exc}")
            _json_response(self, {"error": f"internal error: {exc}"}, 500)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    # Ensure working directories exist
    os.makedirs(SKILL_INPUT_DIR, exist_ok=True)
    os.makedirs(SKILL_OUTPUT_DIR, exist_ok=True)

    server = HTTPServer((HOST, PORT), SkillGatewayHandler)
    print(f"[gateway] OpenClaw Skill Gateway starting on http://{HOST}:{PORT}", flush=True)
    print(f"[gateway] Skills base dir: {SKILLS_BASE}", flush=True)
    print(f"[gateway] Available skills: {', '.join(AVAILABLE_SKILLS)}", flush=True)
    print(f"[gateway] Temp dirs: input={SKILL_INPUT_DIR}, output={SKILL_OUTPUT_DIR}", flush=True)
    print("[gateway] Press Ctrl+C to stop", flush=True)

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n[gateway] Shutting down...", flush=True)
        server.server_close()


if __name__ == "__main__":
    main()
